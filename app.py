# *****************************************************************************
# SECCIÓN 1: NÚCLEO DE NORMALIZACIÓN Y CONVERSIÓN (VERSIÓN EXTENDIDA HH:MM)
# *****************************************************************************
# Esta sección garantiza que los nombres de los deportistas coincidan 
# perfectamente (MatchKey) y que los tiempos de entrenamiento se conviertan 
# a minutos numéricos para poder calcular los KPIs de adherencia (TPI).

import streamlit as st
import pandas as pd
import numpy as np
import re
import io
import zipfile
import unicodedata
import random
from datetime import time, datetime

# Librerías de Word
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -----------------------------------------------------------------------------
# 1.1 NORMALIZACIÓN DE IDENTIDAD (ELIMINA EL ERROR DE NOMBRES CON TILDES)
# -----------------------------------------------------------------------------
def clean_string(text):
    """
    Función de Normalización Absoluta.
    Convierte 'Dagoberto Pérez' y 'DAGOBERTO PEREZ' en el mismo código único.
    Sin esto, el sistema crea filas duplicadas o llena el Maestro con ceros.
    """
    if text is None or pd.isna(text):
        return ""
    
    # Paso 1: Convertir a string, eliminar espacios en los extremos y pasar a MAYÚSCULAS
    nombre_sucio = str(text).strip().upper()
    
    # Paso 2: Normalizar usando NFKD (Descompone caracteres como 'é' en 'e' + '´')
    nombre_normalizado_nfkd = unicodedata.normalize('NFKD', nombre_sucio)
    
    # Paso 3: Filtrar solo los caracteres que no sean marcas de combinación (tildes)
    # Esto elimina físicamente el acento del caracter.
    nombre_limpio_final = "".join([c for c in nombre_normalizado_nfkd if not unicodedata.combining(c)])
    
    return nombre_limpio_final

# -----------------------------------------------------------------------------
# 1.2 CONVERSOR ARITMÉTICO UNIVERSAL (MOTOR ADAPTADO A HH:MM Y TIMEDELTA)
# -----------------------------------------------------------------------------
def to_mins(valor):
    """
    Motor de Conversión de Ingeniería. 
    Transforma cualquier entrada (Texto HH:MM, Timedelta, Decimal) 
    en un número entero de minutos para permitir el cálculo del TPI.
    """
    # Si la celda está vacía, el valor es 0
    if pd.isna(valor):
        return 0
        
    # SALVAVIDAS CRÍTICO 1: Manejo nativo de objetos Timedelta de Pandas
    # Esto evita el fallo de "0 days 02:30:00" que arrojaba ceros en las versiones previas
    if isinstance(valor, pd.Timedelta):
        return int(valor.total_seconds() // 60)
    
    # SALVAVIDAS CRÍTICO 2: El valor es un decimal de Excel (ej: 0.5 equivale a 12 horas)
    # Excel guarda el tiempo como una fracción del día (1.0 = 24h)
    if isinstance(valor, (float, int)):
        if valor == 0:
            return 0
        if valor < 1 and valor > 0:
            return int(round(valor * 1440))
        return int(valor)
        
    # SALVAVIDAS 3: El valor ya es un objeto de tiempo de Python (datetime.time)
    if isinstance(valor, (time, datetime)):
        return (valor.hour * 60) + valor.minute
    
    # Convertir a texto para análisis de patrones
    s = str(valor).strip().lower()
    
    # Lista de exclusión: valores que Strava o Excel ponen cuando no hay actividad
    lista_basura = ['--:--', '0', '', '00:00', '00:00:00', 'nc', 'nan', 'none']
    if s in lista_basura:
        return 0
    
    try:
        # Fallback de seguridad adicional para Timedeltas que llegan como texto
        if 'days' in s or 'day' in s:
            partes_espacio = s.split()
            dias = int(partes_espacio[0])
            tiempo_str = partes_espacio[-1]
            t_parts = tiempo_str.split(':')
            horas = int(t_parts[0])
            minutos = int(t_parts[1].split('.')[0])
            return (dias * 1440) + (horas * 60) + minutos

        # FORMATO PRINCIPAL ESPERADO: HH:MM (Opcionalmente HH:MM:SS)
        if ':' in s:
            partes = s.split(':')
            horas = int(partes[0])
            # Al tomar fijamente el índice [1], funciona perfecto tanto si el archivo
            # trae "02:30" (2 partes) como si trae "02:30:00" (3 partes).
            minutos = int(partes[1].split('.')[0])
            return (horas * 60) + minutos
            
        # SALVAVIDAS 4: Formato de texto de Strava (ej: '1h 22m' o '45min')
        patron_horas = re.search(r'(\d+)\s*h', s)
        patron_minutos = re.search(r'(\d+)\s*m', s)
        
        total_h = int(patron_horas.group(1)) * 60 if patron_horas else 0
        total_m = int(patron_minutos.group(1)) if patron_minutos else 0
        
        if total_h > 0 or total_m > 0:
            return total_h + total_m
            
        # SALVAVIDAS 5: Si por error de tipeo es solo un número suelto (ej: "45")
        if s.isdigit():
            return int(s)
            
    except Exception:
        # Si algo falla en la conversión de esta celda, devolvemos 0 para no romper el lote
        return 0
        
    return 0

# -----------------------------------------------------------------------------
# 1.3 FORMATEADOR VISUAL PARA REPORTES (HH:MM ESTRICTO)
# -----------------------------------------------------------------------------
def to_hhmm_display(minutos):
    """
    Convierte los minutos numéricos de vuelta a un formato legible 
    para las tablas de los reportes Word. Limitado a HH:MM por instrucción.
    """
    horas = int(minutos // 60)
    minutos_restantes = int(minutos % 60)
    
    # Entrega formato 02:30 (sin los segundos finales)
    return f"{horas:02d}:{minutos_restantes:02d}"


# *****************************************************************************
# SECCIÓN 2: MOTORES DE EXTRACCIÓN Y PARSEO (STRAVA Y PLAN INDIVIDUAL)
# *****************************************************************************
# Esta sección lee los archivos Excel de entrada y los transforma en tablas
# estandarizadas. Garantiza que las columnas se identifiquen correctamente
# sin importar ligeras variaciones en los nombres generados por Strava.

# -----------------------------------------------------------------------------
# 2.1 PARSER DE EXCEL SEMANAL (CAPTURADOR DE DATOS REALES DE STRAVA)
# -----------------------------------------------------------------------------
def procesar_strava_excel(archivo_excel):
    """
    Escanea el Excel descargado de Strava buscando las columnas clave.
    Utiliza búsqueda semántica explícita para ser inmune a cambios de formato.
    """
    # Cargar el archivo Excel crudo en un DataFrame de Pandas
    df_raw = pd.read_excel(archivo_excel)
    columnas_originales = df_raw.columns
    
    # Variables de control para almacenar los nombres reales de las columnas encontradas
    columna_nombre = None
    columna_natacion = None
    columna_ciclismo = None
    columna_trote = None
    
    # --- Búsqueda de la columna de Identidad (Deportista / Nombre) ---
    for col in columnas_originales:
        col_limpia = clean_string(str(col))
        if 'DEPORTISTA' in col_limpia or 'NOMBRE' in col_limpia or 'NAME' in col_limpia:
            columna_nombre = col
            break
            
    # Fallback de seguridad extrema: Si no encuentra columna de nombre, asume que es la primera
    if columna_nombre is None:
        columna_nombre = columnas_originales[0]
        
    # --- Búsqueda de la columna de Natación ---
    for col in columnas_originales:
        col_limpia = clean_string(str(col))
        if 'NATAC' in col_limpia or 'PISCINA' in col_limpia or 'SWIM' in col_limpia:
            columna_natacion = col
            break
            
    # --- Búsqueda de la columna de Ciclismo ---
    for col in columnas_originales:
        col_limpia = clean_string(str(col))
        if 'BICI' in col_limpia or 'CICLIS' in col_limpia or 'RIDE' in col_limpia:
            columna_ciclismo = col
            break
            
    # --- Búsqueda de la columna de Trote ---
    for col in columnas_originales:
        col_limpia = clean_string(str(col))
        if 'TROTE' in col_limpia or 'RUN' in col_limpia or 'CARRERA' in col_limpia:
            columna_trote = col
            break

    # --- Construcción del DataFrame Limpio y Purificado ---
    # Aquí creamos la estructura interna que el resto del sistema consumirá sin errores
    df_limpio = pd.DataFrame()
    
    # Asignamos la identidad
    df_limpio['Deportista'] = df_raw[columna_nombre]
    
    # Aplicamos el motor aritmético (to_mins de la Sección 1) a cada celda de las disciplinas
    if columna_natacion is not None:
        df_limpio['N_Mins_Real'] = df_raw[columna_natacion].apply(to_mins)
    else:
        # Si Strava no trae la columna, rellenamos con 0 explícitamente
        df_limpio['N_Mins_Real'] = 0
        
    if columna_ciclismo is not None:
        df_limpio['B_Mins_Real'] = df_raw[columna_ciclismo].apply(to_mins)
    else:
        df_limpio['B_Mins_Real'] = 0
        
    if columna_trote is not None:
        df_limpio['R_Mins_Real'] = df_raw[columna_trote].apply(to_mins)
    else:
        df_limpio['R_Mins_Real'] = 0
        
    # Calculamos el tiempo total real de la semana como la suma exacta de las tres disciplinas
    df_limpio['T_Mins_Real'] = df_limpio['N_Mins_Real'] + df_limpio['B_Mins_Real'] + df_limpio['R_Mins_Real']
    
    return df_limpio

# -----------------------------------------------------------------------------
# 2.2 PARSER DEL PLAN INDIVIDUAL (METAS ESPECÍFICAS)
# -----------------------------------------------------------------------------
def procesar_plan_individual(archivo_plan):
    """
    Carga el archivo Excel con las metas individuales de los deportistas.
    Asegura que la identidad se procese con la misma llave (MatchKey) para un cruce perfecto.
    """
    if archivo_plan is None:
        # Si el usuario decide no cargar un plan individual, devolvemos un DataFrame vacío.
        # El sistema detectará esto y aplicará automáticamente las Metas Globales del Sidebar.
        return pd.DataFrame()
        
    # Cargar el Excel de plan de entrenamiento
    df_plan = pd.read_excel(archivo_plan)
    
    # Verificación de seguridad: Evitar procesar archivos que vengan en blanco
    if df_plan.empty:
        return pd.DataFrame()
    
    # Asumimos estrictamente que la primera columna contiene la identidad del deportista
    columna_nombres_plan = df_plan.columns[0]
    
    # Generamos la llave maestra de cruce usando la normalización de la Sección 1
    df_plan['MatchKey'] = df_plan[columna_nombres_plan].apply(clean_string)
    
    return df_plan
    
# *****************************************************************************
# SECCIÓN 3: MOTOR NARRATIVO PRO CHILE - NIVEL 2 (CONTEXTUAL)
# *****************************************************************************

import random

PILAS_COMENTARIOS = {}

# -----------------------------------------------------------------------------
# 3.1 CONSTANTES Y CLASIFICADORES
# -----------------------------------------------------------------------------

UMBRAL_MEJORA = 15
UMBRAL_BAJA = -15
ZONA_ROJO = 70
ZONA_AMARILLO = 90

def clasificar_estado(diff_hist_mins):
    if diff_hist_mins is None:
        return 'primer_registro'
    if diff_hist_mins > UMBRAL_MEJORA:
        return 'mejora'
    if diff_hist_mins < UMBRAL_BAJA:
        return 'baja'
    return 'estable'

def clasificar_zona_tpi(tpi_valor):
    if tpi_valor < ZONA_ROJO:
        return 'rojo'
    if tpi_valor <= ZONA_AMARILLO:
        return 'amarillo'
    return 'verde'

def clasificar_contexto_disciplinas(mins_n, mins_b, mins_r):
    faltantes = []
    if mins_n == 0: faltantes.append('natación')
    if mins_b == 0: faltantes.append('ciclismo')
    if mins_r == 0: faltantes.append('trote')
    if not faltantes:
        return 'completo', []
    if len(faltantes) == 3:
        return 'sin_actividad', faltantes
    return 'incompleto', faltantes

def obtener_frase_base(clave, pool_frases):
    global PILAS_COMENTARIOS
    if clave not in PILAS_COMENTARIOS or not PILAS_COMENTARIOS[clave]:
        temp = [str(f) for f in pool_frases]
        random.shuffle(temp)
        PILAS_COMENTARIOS[clave] = temp
    return PILAS_COMENTARIOS[clave].pop()

# -----------------------------------------------------------------------------
# 3.2 BANCOS DE FRASES CONTEXTUALES
# -----------------------------------------------------------------------------

FRASES = {

    # --- ESTADO DEL ATLETA (tono 3ª persona - reporte grupal) ---
    'estado_mejora_grupal': [
        "{atleta} llega con más energía que la semana pasada. La diferencia se nota en los números.",
        "Semana ascendente para {atleta}. Su media histórica quedó atrás y eso es exactamente el objetivo.",
        "{atleta} entregó más de lo habitual. El cuerpo respondió y el registro lo confirma.",
        "Evolución concreta de {atleta}. No es opinión, son los minutos acumulados los que hablan.",
        "La tendencia de {atleta} va hacia arriba. Una semana para recordar en su historial personal.",
        "Superó su propia media histórica. {atleta} demuestra que el techo siempre puede subir un poco más.",
        "{atleta} se exigió más que de costumbre y el resultado está a la vista en la tabla.",
        "Semana destacada. {atleta} rompió su promedio habitual con una carga que merece reconocimiento.",
    ],
    'estado_estable_grupal': [
        "{atleta} mantiene su ritmo habitual. La consistencia también es una forma de progreso.",
        "Semana dentro de lo esperado para {atleta}. Sostener el nivel no es poca cosa.",
        "{atleta} en su línea. Sin grandes cambios, pero sin retrocesos. Eso tiene valor.",
        "Regularidad de {atleta}. El entrenamiento invisible se construye con semanas así.",
        "{atleta} no sorprende, pero tampoco decepciona. La constancia es su firma.",
        "Semana estable para {atleta}. El plan se respeta y eso es el primer paso.",
        "{atleta} en crucero. Volumen similar a su media histórica, lo que refleja un buen control de carga.",
        "Sin grandes oscilaciones. {atleta} gestiona su semana con madurez deportiva.",
    ],
    'estado_baja_grupal': [
        "{atleta} tuvo una semana más liviana que de costumbre. Puede ser estrategia o puede ser señal.",
        "El volumen de {atleta} bajó respecto a su media histórica. Una semana para retomar el ritmo.",
        "Semana de menor carga para {atleta}. El descanso tiene su lugar, pero hay que volver pronto.",
        "{atleta} registró menos minutos que su promedio. La próxima semana es la oportunidad de revertirlo.",
        "Baja en el volumen de {atleta}. No es drama, pero sí una señal para revisar la semana.",
        "{atleta} estuvo por debajo de su media. A veces el cuerpo manda, otras veces manda la agenda.",
        "Semana corta para {atleta}. El historial personal pide un poco más la próxima vez.",
        "El registro de {atleta} estuvo bajo su nivel habitual. La constancia se recupera, siempre.",
    ],
    'estado_primer_registro_grupal': [
        "{atleta} estrena su historial en MetriKM. Bienvenido al seguimiento semanal.",
        "Primera semana registrada para {atleta}. Desde aquí se construye la base de datos personal.",
        "{atleta} entra al sistema. Este es el punto de partida desde el cual todo se mide.",
        "Debut en el registro semanal. {atleta} comienza a escribir su historial deportivo en el club.",
        "Primera huella de {atleta} en el Maestro. Ahora sí hay datos con qué comparar las próximas semanas.",
    ],

    # --- ESTADO DEL ATLETA (tono 2ª persona - ficha individual) ---
    'estado_mejora_individual': [
        "Esta semana superaste tu propia media histórica. Eso no pasa solo, es el resultado del trabajo.",
        "Tus números de esta semana están por encima de tu promedio habitual. Sigue en esa dirección.",
        "Rompiste tu media histórica. Pequeño salto, gran señal de progresión.",
        "Mejor semana que tu promedio. El esfuerzo se tradujo directamente en los registros.",
        "Esta semana entrenaste más de lo que acostumbras. Tu historial personal te lo va a agradecer.",
        "Superaste tu línea base personal. Así se construye la progresión semana a semana.",
    ],
    'estado_estable_individual': [
        "Semana dentro de tu rango habitual. Mantener el nivel también es avanzar.",
        "Tus números están alineados con tu media histórica. Consistencia que se nota.",
        "Sin grandes cambios respecto a tus semanas anteriores. La regularidad es tu mejor aliada.",
        "Estuviste en tu zona de confort de volumen. A veces consolidar es lo correcto.",
        "Semana estable. Tu historial muestra que este es tu ritmo base, y eso tiene valor.",
        "Dentro de tu media histórica. El cuerpo y el plan estuvieron sincronizados esta semana.",
    ],
    'estado_baja_individual': [
        "Esta semana entrenaste menos que tu promedio habitual. La próxima es la oportunidad de retomar.",
        "Tus minutos bajaron respecto a tu media histórica. ¿Semana complicada o descanso planificado?",
        "Registro por debajo de tu línea base. No es para alarmarse, pero sí para retomar el ritmo pronto.",
        "Volumen más bajo que de costumbre. Tu historial pide un poco más la próxima semana.",
        "Estuviste bajo tu media personal. El cuerpo a veces necesita bajar, pero hay que volver.",
        "Semana liviana comparada con tu historial. La consistencia se recupera, una semana a la vez.",
    ],
    'estado_primer_registro_individual': [
        "Esta es tu primera semana registrada en MetriKM. Desde aquí se construye todo.",
        "Bienvenido al seguimiento semanal. Este registro es el punto de partida de tu historial personal.",
        "Primera semana en el sistema. La próxima ya tendrás con qué compararte.",
        "Debut en el Maestro. A partir de ahora cada semana suma a tu base de datos personal.",
    ],

    # --- ZONA TPI (tono grupal) ---
    'tpi_rojo_grupal': [
        "El plan existía. Los minutos, menos. {atleta} tiene una conversación pendiente con su agenda.",
        "Adherencia baja para {atleta} esta semana. El plan es una intención, el registro es la realidad.",
        "{atleta} cumplió una parte del plan. La otra parte queda pendiente para la próxima semana.",
        "TPI en zona roja para {atleta}. No es el fin del mundo, pero sí una señal clara.",
        "El plan de {atleta} y su semana real tuvieron poco contacto. Hay margen de mejora evidente.",
        "Cumplimiento bajo esta semana para {atleta}. La planificación vale lo que se ejecuta.",
    ],
    'tpi_amarillo_grupal': [
        "{atleta} cumplió una buena parte del plan. Faltó poco para llegar a la zona verde.",
        "Adherencia razonable de {atleta}. El trabajo estuvo, aunque no al 100% de lo planificado.",
        "{atleta} en zona amarilla. Bien encaminado, con espacio para ajustar la próxima semana.",
        "Cerca del objetivo de cumplimiento. {atleta} tiene el motor encendido, le falta un poco de bencina.",
        "Semana de buena adherencia para {atleta}, sin ser perfecta. Eso también es parte del proceso.",
        "{atleta} rozó la zona verde. Un pequeño ajuste la próxima semana puede marcar la diferencia.",
    ],
    'tpi_verde_grupal': [
        "{atleta} ejecutó el plan con precisión. Eso no es suerte, es disciplina.",
        "Zona verde de adherencia para {atleta}. El plan se respetó y los números lo demuestran.",
        "Cumplimiento ejemplar de {atleta}. Cuando el plan y la ejecución coinciden, el progreso es inevitable.",
        "{atleta} clavó el plan esta semana. Referente de adherencia para el club.",
        "TPI en verde para {atleta}. Planificó bien y ejecutó mejor. Así se hace.",
        "Adherencia total de {atleta}. El equipo técnico tiene poco que corregir esta semana.",
    ],

    # --- ZONA TPI (tono individual) ---
    'tpi_rojo_individual': [
        "Tu adherencia al plan estuvo baja esta semana. El plan es una hoja de ruta, no una sugerencia.",
        "TPI en zona roja. Hubo una brecha importante entre lo planificado y lo ejecutado.",
        "Esta semana el plan y tu semana real no se encontraron mucho. La próxima es la oportunidad.",
        "Cumplimiento bajo. ¿Qué pasó esta semana? Vale la pena revisarlo antes de la siguiente.",
        "Tu adherencia necesita atención. El plan existe por una razón, y esta semana quedó corto.",
    ],
    'tpi_amarillo_individual': [
        "Buena adherencia, aunque con espacio para mejorar. Estuviste cerca de la zona verde.",
        "Tu TPI está en amarillo. El esfuerzo estuvo, faltó un poco más para cerrar el plan completo.",
        "Cerca del objetivo de cumplimiento. Un ajuste pequeño la próxima semana puede llevarte al verde.",
        "Adherencia razonable. El plan se respetó en buena parte, pero hay margen de mejora.",
        "Zona amarilla de TPI. Bien encaminado, sin llegar al tope. La próxima semana tienes el dato.",
    ],
    'tpi_verde_individual': [
        "Adherencia en zona verde. Planificaste bien y ejecutaste mejor. Eso es lo que construye progreso.",
        "TPI en verde. Esta semana el plan y tu ejecución estuvieron perfectamente alineados.",
        "Cumplimiento ejemplar. Pocas cosas son más valiosas en el entrenamiento que respetar el plan.",
        "Zona verde de adherencia. El equipo técnico no tiene mucho que corregirte esta semana.",
        "Ejecutaste el plan con precisión. Eso no es casualidad, es disciplina sostenida.",
    ],

    # --- CONTEXTO DISCIPLINAS (grupal) ---
    'completo_grupal': [
        "{atleta} tocó el agua, los pedales y el asfalto. Triatleta completo en toda la extensión de la palabra.",
        "Las tres disciplinas registradas para {atleta}. Eso es lo que diferencia a un triatleta de un deportista de una sola disciplina.",
        "{atleta} no dejó ningún frente abandonado esta semana. Natación, ciclismo y trote, todos presentes.",
        "Semana completa para {atleta}. Tres disciplinas, tres registros, cero excusas.",
        "{atleta} cumplió el requisito fundamental: las tres disciplinas activas. El resto son detalles.",
    ],
    'completo_individual': [
        "Tres disciplinas registradas. Eso es exactamente lo que define a un triatleta.",
        "Natación, ciclismo y trote en tu registro semanal. Semana completa en toda regla.",
        "Las tres disciplinas activas esta semana. Sin atajos, sin disciplinas olvidadas.",
        "Tocaste el agua, los pedales y el asfalto. Semana de triatleta completo.",
        "Tres de tres. Así se construye la base para competir en las tres disciplinas.",
    ],
    'incompleto_grupal': [
        "{atleta} dejó {faltantes} fuera del registro esta semana. El triatlón tiene tres partes, no dos.",
        "Faltó {faltantes} en la semana de {atleta}. Una disciplina ausente es una debilidad que se acumula.",
        "{atleta} entrenó bien, pero {faltantes} quedó en el tintero. La próxima semana, que no falte.",
        "Semana incompleta para {atleta}. {faltantes} no aparece en el registro y eso tiene un costo.",
        "{atleta} y {faltantes} no se vieron esta semana. Habrá que reencontrarse pronto.",
    ],
    'incompleto_individual': [
        "Esta semana te faltó {faltantes}. El triatlón no perdona las disciplinas abandonadas.",
        "Faltó {faltantes} en tu registro. Una semana sin esa disciplina es una semana de ventaja para quienes sí la hicieron.",
        "Tu registro quedó incompleto: {faltantes} no aparece esta semana. Ojo con eso.",
        "Sin {faltantes} esta semana. La próxima, que no falte ninguna disciplina.",
        "{faltantes} ausente en tu semana. Recuerda que el triatlón cobra las tres disciplinas el día de la carrera.",
    ],
    'sin_actividad_grupal': [
        "{atleta} no registró actividad esta semana. Descanso, viaje o contratiempo, la próxima semana cuenta.",
        "Semana en blanco para {atleta}. El cuerpo a veces necesita parar, pero el plan no espera.",
        "{atleta} no apareció en el registro esta semana. Esperamos verlo de vuelta pronto.",
    ],
    'sin_actividad_individual': [
        "Sin actividad registrada esta semana. La próxima semana es la oportunidad de retomar.",
        "No hubo registro esta semana. Pase lo que haya pasado, la siguiente semana empieza desde cero.",
        "Semana en blanco en tu historial. El plan sigue esperando, sin juzgar.",
    ],

    # --- DISCIPLINAS ESPECÍFICAS (grupal) ---
    'natacion_grupal': [
        "{atleta} sumó {tiempo} en la piscina. El agua no miente y este registro lo confirma.",
        "Volumen de natación sólido para {atleta}: {tiempo}. La base acuática se construye con semanas así.",
        "{atleta} y la piscina tuvieron una buena semana juntos: {tiempo} de trabajo acuático.",
        "{tiempo} de natación para {atleta}. Cada largo suma a la base aeróbica que se necesita en carrera.",
        "Registro acuático de {atleta}: {tiempo}. La piscina es donde se gana el tiempo en el segmento más técnico.",
        "{atleta} acumuló {tiempo} en el agua. Eso es fondo, técnica y resistencia en un solo número.",
    ],
    'ciclismo_grupal': [
        "{atleta} puso {tiempo} sobre los pedales esta semana. El segmento más largo del triatlón bien cubierto.",
        "{tiempo} de ciclismo para {atleta}. Las piernas tienen memoria y este volumen les habla.",
        "{atleta} y la bicicleta: {tiempo} juntos esta semana. El motor aeróbico agradece cada pedalada.",
        "Rodaje de {tiempo} para {atleta}. En ciclismo, el volumen es la base de todo lo demás.",
        "{atleta} acumuló {tiempo} en bicicleta. Eso es inversión directa en el segmento más largo de la carrera.",
        "{tiempo} de pedaleo para {atleta}. Las watts no se improvisan, se construyen semana a semana.",
    ],
    'trote_grupal': [
        "{atleta} cerró con {tiempo} de trote. La carrera a pie es donde se define el triatlón y este volumen lo sabe.",
        "{tiempo} de running para {atleta}. El asfalto tiene sus propias reglas y {atleta} las respeta.",
        "{atleta} sumó {tiempo} de trote esta semana. Las piernas cansadas de la bici necesitan este trabajo.",
        "Registro de carrera a pie de {atleta}: {tiempo}. El último segmento se gana entrenando este volumen.",
        "{tiempo} de zancada para {atleta}. El trote es donde muchos triatletas pierden o ganan la carrera.",
        "{atleta} y el asfalto: {tiempo} esta semana. Consistencia en running es consistencia en resultados.",
    ],
    'natacion_individual': [
        "Sumaste {tiempo} en la piscina. El agua es el segmento más técnico y este volumen construye base.",
        "{tiempo} de natación en tu registro. Cada sesión en el agua mejora algo que no se ve en la bici ni en el trote.",
        "Tu volumen acuático esta semana: {tiempo}. La piscina te devuelve lo que le das, sin excepciones.",
        "{tiempo} en el agua. Eso es técnica, resistencia y confianza para el segmento de apertura.",
        "Natación: {tiempo} esta semana. El primer segmento del triatlón se gana aquí, en el entrenamiento.",
    ],
    'ciclismo_individual': [
        "{tiempo} sobre los pedales esta semana. El segmento más largo del triatlón bien trabajado.",
        "Tu volumen de ciclismo: {tiempo}. Las piernas tienen memoria y este trabajo suma.",
        "{tiempo} de rodaje en tu registro. La bicicleta es donde se construye el motor aeróbico.",
        "Ciclismo: {tiempo} esta semana. Ese volumen se traduce directamente en capacidad para el resto del recorrido.",
        "{tiempo} de pedaleo. En el triatlón, la bici es el segmento que más tiempo consume y más base requiere.",
    ],
    'trote_individual': [
        "{tiempo} de trote en tu registro. La carrera a pie es donde el triatlón se decide.",
        "Tu volumen de running esta semana: {tiempo}. Las piernas cansadas de la bici necesitan este trabajo.",
        "{tiempo} de zancada. El último segmento se entrena así, semana a semana.",
        "Trote: {tiempo} esta semana. El asfalto tiene sus propias reglas y este volumen las respeta.",
        "{tiempo} de carrera a pie. Aquí es donde muchos triatlones se ganan o se pierden.",
    ],
}

# -----------------------------------------------------------------------------
# 3.3 FUNCIÓN PRINCIPAL DE GENERACIÓN DE COMENTARIOS
# -----------------------------------------------------------------------------

def generar_comentario(datos_de_fila, nombre_categoria, rank_posicion,
                       diff_hist_mins=None, destino='grupal'):
    """
    Motor contextual Nivel 2.
    Genera comentarios según estado del atleta, zona TPI, 
    contexto de disciplinas y destino del reporte.
    """
    atleta = str(datos_de_fila.get('Deportista', 'Atleta TYM'))
    tpi_global = datos_de_fila.get('TPI_Global', 0)
    mins_n = datos_de_fila.get('N_Mins_Real', 0)
    mins_b = datos_de_fila.get('B_Mins_Real', 0)
    mins_r = datos_de_fila.get('R_Mins_Real', 0)

    # Tiempo formateado según disciplina
    if nombre_categoria == 'Natación':
        tiempo = to_hhmm_display(mins_n)
    elif nombre_categoria == 'Bicicleta':
        tiempo = to_hhmm_display(mins_b)
    elif nombre_categoria == 'Trote':
        tiempo = to_hhmm_display(mins_r)
    else:
        tiempo = to_hhmm_display(datos_de_fila.get('T_Mins_Real', 0))

    # Clasificadores
    estado = clasificar_estado(diff_hist_mins)
    zona = clasificar_zona_tpi(tpi_global)
    contexto_disc, faltantes = clasificar_contexto_disciplinas(mins_n, mins_b, mins_r)
    faltantes_str = ' y '.join(faltantes) if faltantes else ''

    comentarios = []

    # --- BLOQUE 1: ESTADO DEL ATLETA ---
    clave_estado = f'estado_{estado}_{destino}'
    pool_estado = FRASES.get(clave_estado, FRASES.get(f'estado_{estado}_grupal', []))
    if pool_estado:
        frase = obtener_frase_base(clave_estado, pool_estado)
        comentarios.append(frase.replace('{atleta}', atleta).replace('{tiempo}', tiempo))

    # --- BLOQUE 2: ZONA TPI ---
    clave_tpi = f'tpi_{zona}_{destino}'
    pool_tpi = FRASES.get(clave_tpi, FRASES.get(f'tpi_{zona}_grupal', []))
    if pool_tpi:
        frase = obtener_frase_base(clave_tpi, pool_tpi)
        comentarios.append(frase.replace('{atleta}', atleta).replace('{tiempo}', tiempo))

    # --- BLOQUE 3: CONTEXTO DISCIPLINAS ---
    clave_disc = f'{contexto_disc}_{destino}'
    pool_disc = FRASES.get(clave_disc, FRASES.get(f'{contexto_disc}_grupal', []))
    if pool_disc:
        frase = obtener_frase_base(clave_disc, pool_disc)
        comentarios.append(
            frase.replace('{atleta}', atleta)
                 .replace('{tiempo}', tiempo)
                 .replace('{faltantes}', faltantes_str)
        )

    # --- BLOQUE 4: DISCIPLINA ESPECÍFICA (si aplica) ---
    mapa_disc = {
        'Natación': f'natacion_{destino}',
        'Bicicleta': f'ciclismo_{destino}',
        'Trote': f'trote_{destino}'
    }
    if nombre_categoria in mapa_disc:
        clave_especifica = mapa_disc[nombre_categoria]
        pool_esp = FRASES.get(clave_especifica, [])
        if pool_esp:
            frase = obtener_frase_base(clave_especifica, pool_esp)
            comentarios.append(frase.replace('{atleta}', atleta).replace('{tiempo}', tiempo))

    # Unir los bloques en un párrafo coherente
    comentario_final = ' '.join(comentarios)

    # Bonus de liderazgo para el 1° lugar
    if rank_posicion == 1 and destino == 'grupal':
        comentario_final = f"🏆 {comentario_final}"

    return comentario_final if comentario_final.strip() else f"{atleta} registró actividad esta semana."
    
# *****************************************************************************
# SECCIÓN 4: MOTOR DE CÁLCULO DE ADHERENCIA (TPI - REGLA 4.3)
# *****************************************************************************
# Esta sección es el corazón analítico del sistema. Cruza los datos reales 
# obtenidos de Strava con las metas individuales o globales para calcular el TPI
# (Índice de Rendimiento TYM). Garantiza que no existan valores nulos (0.0%)
# ni errores de división por cero. Ahora implementa un TPI Global Ponderado.

def calcular_kpis_tym(df_real, df_plan, metas_globales):
    """
    Calcula la adherencia al entrenamiento (TPI).
    Implementa la 'Cascada de Metas': Si un atleta no tiene un plan individual
    cargado en el Excel, el sistema utilizará automáticamente las metas globales
    definidas en el panel lateral (Sidebar).
    """
    # 1. Asegurar la existencia de la llave de cruce en los datos reales
    df_real['MatchKey'] = df_real['Deportista'].apply(clean_string)
    
    # 2. Cruce de datos (Merge) con el Plan Individual
    if df_plan is not None and not df_plan.empty:
        # Aseguramos que el plan también tenga su llave de cruce
        # Asumimos que la columna 0 del plan es el nombre del deportista
        nombre_col_plan = df_plan.columns[0]
        df_plan['MatchKey'] = df_plan[nombre_col_plan].apply(clean_string)
        
        # Realizamos un LEFT JOIN para mantener a todos los deportistas reales,
        # incluso si no tienen un plan individual cargado en su fila.
        df_merged = pd.merge(df_real, df_plan, on='MatchKey', how='left', suffixes=('', '_plan'))
    else:
        # Si no se subió el archivo de plan individual, trabajamos directamente con los datos reales
        df_merged = df_real.copy()

    # 3. Función interna para procesar la aritmética fila por fila (Deportista por Deportista)
    def aplicar_logica_tpi_fila(row):
        resultados_kpi = {}
        
        # ---------------------------------------------------------------------
        # --- A. EVALUACIÓN DE NATACIÓN ---
        # ---------------------------------------------------------------------
        # Obtener meta de horas (Cascada: Plan Individual -> Meta Global)
        horas_meta_natacion = row.get('Natacion_Hrs_plan')
        if pd.isna(horas_meta_natacion):
            horas_meta_natacion = metas_globales['N_H']
            
        # Obtener meta de sesiones (Cascada: Plan Individual -> Meta Global)
        sesiones_meta_natacion = row.get('Natacion_Ses_plan')
        if pd.isna(sesiones_meta_natacion):
            sesiones_meta_natacion = metas_globales['N_S']
            
        # Obtener el dato real procesado en la Sección 2
        minutos_reales_natacion = row.get('N_Mins_Real', 0)
        
        # Cálculo de Índices para Natación
        # VCI (Volume Compliance Index): Porcentaje de horas cumplidas respecto al plan
        if horas_meta_natacion > 0:
            vci_natacion = (minutos_reales_natacion / (horas_meta_natacion * 60)) * 100
        else:
            vci_natacion = 0
            
        # SEI (Session Execution Index): Cumplimiento de sesiones
        # Si el atleta entrenó más de 0 minutos y la meta es mayor a 0, asimilamos el 100% de la sesión
        if minutos_reales_natacion > 0 and sesiones_meta_natacion > 0:
            sei_natacion = (100 / sesiones_meta_natacion)
        else:
            sei_natacion = 0
        
        # TPI Natación (Regla temporal: 100% Volumen hasta tener sesiones), tope máximo de 115% de cumplimiento
        tpi_natacion_crudo = vci_natacion
        resultados_kpi['TPI_Natacion'] = min(tpi_natacion_crudo, 115)
        resultados_kpi['Natacion_Plan_Hrs'] = horas_meta_natacion

        # ---------------------------------------------------------------------
        # --- B. EVALUACIÓN DE CICLISMO ---
        # ---------------------------------------------------------------------
        # Obtener meta de horas
        horas_meta_ciclismo = row.get('Ciclismo_Hrs_plan')
        if pd.isna(horas_meta_ciclismo):
            horas_meta_ciclismo = metas_globales['B_H']
            
        # Obtener meta de sesiones
        sesiones_meta_ciclismo = row.get('Ciclismo_Ses_plan')
        if pd.isna(sesiones_meta_ciclismo):
            sesiones_meta_ciclismo = metas_globales['B_S']
            
        # Obtener el dato real procesado
        minutos_reales_ciclismo = row.get('B_Mins_Real', 0)
        
        # Cálculo de Índices para Ciclismo
        if horas_meta_ciclismo > 0:
            vci_ciclismo = (minutos_reales_ciclismo / (horas_meta_ciclismo * 60)) * 100
        else:
            vci_ciclismo = 0
            
        if minutos_reales_ciclismo > 0 and sesiones_meta_ciclismo > 0:
            sei_ciclismo = (100 / sesiones_meta_ciclismo)
        else:
            sei_ciclismo = 0
        
        # TPI Ciclismo (Temporal: 100% Volumen)
        tpi_ciclismo_crudo = vci_ciclismo
        resultados_kpi['TPI_Ciclismo'] = min(tpi_ciclismo_crudo, 115)
        resultados_kpi['Ciclismo_Plan_Hrs'] = horas_meta_ciclismo

        # ---------------------------------------------------------------------
        # --- C. EVALUACIÓN DE TROTE ---
        # ---------------------------------------------------------------------
        # Obtener meta de horas
        horas_meta_trote = row.get('Trote_Hrs_plan')
        if pd.isna(horas_meta_trote):
            horas_meta_trote = metas_globales['T_H']
            
        # Obtener meta de sesiones
        sesiones_meta_trote = row.get('Trote_Ses_plan')
        if pd.isna(sesiones_meta_trote):
            sesiones_meta_trote = metas_globales['T_S']
            
        # Obtener el dato real procesado
        minutos_reales_trote = row.get('R_Mins_Real', 0)
        
        # Cálculo de Índices para Trote
        if horas_meta_trote > 0:
            vci_trote = (minutos_reales_trote / (horas_meta_trote * 60)) * 100
        else:
            vci_trote = 0
            
        if minutos_reales_trote > 0 and sesiones_meta_trote > 0:
            sei_trote = (100 / sesiones_meta_trote)
        else:
            sei_trote = 0
        
        # TPI Trote (Temporal: 100% Volumen)
        tpi_trote_crudo = vci_trote
        resultados_kpi['TPI_Trote'] = min(tpi_trote_crudo, 115)
        resultados_kpi['Trote_Plan_Hrs'] = horas_meta_trote

        # ---------------------------------------------------------------------
        # --- D. CÁLCULO DE INDICADORES GLOBALES DEL DEPORTISTA ---
        # ---------------------------------------------------------------------
        # CÁLCULO DEL TPI GLOBAL PONDERADO
        # Calculamos el peso relativo de cada disciplina dentro del plan total semanal
        total_horas_plan = horas_meta_natacion + horas_meta_ciclismo + horas_meta_trote
        
        if total_horas_plan > 0:
            peso_n = horas_meta_natacion / total_horas_plan
            peso_b = horas_meta_ciclismo / total_horas_plan
            peso_r = horas_meta_trote / total_horas_plan
            
            # El TPI Global ahora es la suma de los éxitos multiplicados por el peso de su disciplina
            resultados_kpi['TPI_Global'] = (
                (resultados_kpi['TPI_Natacion'] * peso_n) +
                (resultados_kpi['TPI_Ciclismo'] * peso_b) +
                (resultados_kpi['TPI_Trote'] * peso_r)
            )
        else:
            # Fallback de seguridad en caso de que el plan total sea 0 horas
            resultados_kpi['TPI_Global'] = 0.0
        
        # Validación Estricta de "Triatleta Completo"
        # Debe registrar estrictamente más de 0 minutos en las TRES disciplinas.
        # Esto soluciona de raíz el fallo del Word donde reportaba "0 Triatletas Completos".
        if (minutos_reales_natacion > 0) and (minutos_reales_ciclismo > 0) and (minutos_reales_trote > 0):
            resultados_kpi['Es_Completo'] = True
        else:
            resultados_kpi['Es_Completo'] = False
        
        # Retornamos los cálculos transformados en una Serie de Pandas
        return pd.Series(resultados_kpi)

    # 4. Ejecución del motor: Aplicamos la lógica fila por fila a todos los deportistas
    df_resultados_kpi = df_merged.apply(aplicar_logica_tpi_fila, axis=1)
    
    # 5. Concatenación: Unimos los datos reales (Strava) con los KPI recién calculados
    df_final_procesado = pd.concat([df_merged, df_resultados_kpi], axis=1)
    
    return df_final_procesado
    


# *****************************************************************************
# SECCIÓN 5: MOTOR DE PERSISTENCIA Y ACTUALIZACIÓN DEL MAESTRO (FASE 1 - EXPANDIDA)
# *****************************************************************************
# Esta sección actualiza el libro Excel Histórico preservando las semanas
# anteriores. Implementa un blindaje contra columnas duplicadas (sufijos _x, _y),
# aplica la conversión a formato HH:MM estricto, reordena las columnas de 
# semanas, inyecta la hoja de KPIs, redondea el CV a 2 decimales y aplica
# un ordenamiento estricto y jerárquico a las pestañas del archivo final.
# NUEVO: Genera y actualiza automáticamente la hoja 'Histórico_TPI'.

def actualizar_maestro_tym(dict_dfs_originales, df_semana_actual, etiqueta_semana):
    """
    Actualiza hoja por hoja el archivo Maestro.
    Garantiza que la información se indexe correctamente a cada atleta
    y que la historia no se borre ni se convierta en ceros, formateando todo a HH:MM.
    """
    dict_dfs_actualizados = {}

    # --- NUEVA EXPANSIÓN: ASEGURAR EXISTENCIA DE LA HOJA HISTÓRICO TPI ---
    # Verificamos si ya existe una hoja para el histórico de adherencia
    hoja_tpi_existente = any(clean_string(k) in ['HISTORICO_TPI', 'HISTÓRICO_TPI', 'HISTORICOTPI'] for k in dict_dfs_originales.keys())
    
    if not hoja_tpi_existente:
        # Si no existe, la creamos como una hoja base con los deportistas de la semana actual
        df_historico_tpi = pd.DataFrame()
        df_historico_tpi['Deportista'] = df_semana_actual['Deportista']
        dict_dfs_originales['Histórico_TPI'] = df_historico_tpi

    # 1. Mapeo Extendido de Disciplinas (Sincronización de Pestañas)
    mapeo_hojas_a_datos = {
        'TIEMPO TOTAL': 'T_Mins_Real',
        'NATACION': 'N_Mins_Real',
        'NATACIÓN': 'N_Mins_Real',
        'BICICLETA': 'B_Mins_Real',
        'CICLISMO': 'B_Mins_Real',
        'TROTE': 'R_Mins_Real',
        'RUNNING': 'R_Mins_Real',
        'CV': 'TPI_Global',
        'HISTORICO_TPI': 'TPI_Global',
        'HISTÓRICO_TPI': 'TPI_Global',
        'HISTORICOTPI': 'TPI_Global'
    }

    # Normalizamos los nombres de las pestañas reales del archivo subido
    hojas_reales_normalizadas = {clean_string(nombre): nombre for nombre in dict_dfs_originales.keys()}

    # Diccionario para mapear MatchKey a Nombres reales de la semana actual (Para atletas nuevos)
    mapeo_nombres_nuevos = df_semana_actual.set_index('MatchKey')['Deportista'].to_dict()

    # 2. Bucle de Procesamiento: Hoja por Hoja
    for nombre_hoja_original in dict_dfs_originales.keys():
        nombre_normalizado = clean_string(nombre_hoja_original)

        # Verificamos si la pestaña actual es una disciplina a actualizar
        if nombre_normalizado in mapeo_hojas_a_datos:
            columna_datos_a_extraer = mapeo_hojas_a_datos[nombre_normalizado]
            df_hoja_historia = dict_dfs_originales[nombre_hoja_original].copy()

            # --- IDENTIFICACIÓN DEL DEPORTISTA ---
            col_identidad_maestro = 'Nombre' if 'Nombre' in df_hoja_historia.columns else \
                                   ('Deportista' if 'Deportista' in df_hoja_historia.columns else df_hoja_historia.columns[0])

            # Generamos la llave maestra (MatchKey) para un cruce seguro
            df_hoja_historia['MatchKey'] = df_hoja_historia[col_identidad_maestro].apply(clean_string)

            # --- BLINDAJE ANTI-DUPLICADOS Y LIMPIEZA DE BASURA ---
            columnas_basura = [c for c in df_hoja_historia.columns if etiqueta_semana in str(c) or str(c).endswith('_x') or str(c).endswith('_y')]
            if columnas_basura:
                df_hoja_historia = df_hoja_historia.drop(columns=columnas_basura)

            # --- SALVACIÓN Y PURIFICACIÓN DEL HISTORIAL ---
            columnas_semanas_viejas = [c for c in df_hoja_historia.columns if str(c).startswith('Sem')]

            for col_vieja in columnas_semanas_viejas:
                if columna_datos_a_extraer != 'TPI_Global':
                    df_hoja_historia[col_vieja] = df_hoja_historia[col_vieja].apply(to_mins)
                else:
                    # En las hojas de TPI/CV forzamos a numérico y aplicamos el límite de 2 decimales
                    df_hoja_historia[col_vieja] = pd.to_numeric(df_hoja_historia[col_vieja], errors='coerce').fillna(0.0).round(2)

            # --- PREPARACIÓN DE LA NOVEDAD (DATOS DE LA SEMANA) ---
            df_novedad = df_semana_actual[['MatchKey', columna_datos_a_extraer]].copy()

            if columna_datos_a_extraer != 'TPI_Global':
                df_novedad[etiqueta_semana] = df_novedad[columna_datos_a_extraer]
            else:
                # La novedad de TPI/CV también se redondea a 2 decimales
                df_novedad[etiqueta_semana] = (df_novedad[columna_datos_a_extraer] / 100.0).round(2)

            # Eliminamos duplicados
            df_novedad = df_novedad.drop_duplicates(subset=['MatchKey'], keep='first')

            # --- UNIÓN DEL HISTÓRICO CON LA NOVEDAD ---
            df_actualizado = pd.merge(
                df_hoja_historia,
                df_novedad[['MatchKey', etiqueta_semana]],
                on='MatchKey',
                how='outer'
            )

            df_actualizado[etiqueta_semana] = df_actualizado[etiqueta_semana].fillna(0.0)

            # --- GESTIÓN DE ATLETAS NUEVOS ---
            mask_nuevos = df_actualizado[col_identidad_maestro].isna()
            df_actualizado.loc[mask_nuevos, col_identidad_maestro] = df_actualizado.loc[mask_nuevos, 'MatchKey'].map(mapeo_nombres_nuevos)

            for col_vieja in columnas_semanas_viejas:
                df_actualizado[col_vieja] = df_actualizado[col_vieja].fillna(0.0)

            # --- ORDENAMIENTO DE COLUMNAS DE SEMANAS ---
            columnas_fijas = [c for c in df_actualizado.columns if not str(c).startswith('Sem')]
            columnas_temporales = sorted([c for c in df_actualizado.columns if str(c).startswith('Sem')])
            df_actualizado = df_actualizado[columnas_fijas + columnas_temporales]

            # --- RECALCULO DE MÉTRICAS (TIEMPO ACUMULADO) ---
            columnas_todas_semanas = [c for c in df_actualizado.columns if str(c).startswith('Sem')]

            if columna_datos_a_extraer != 'TPI_Global':
                if 'Tiempo Acumulado' in df_actualizado.columns:
                    df_actualizado['Tiempo Acumulado'] = df_actualizado[columnas_todas_semanas].sum(axis=1)

                if 'Promedio' in df_actualizado.columns:
                    df_actualizado['Promedio'] = df_actualizado[columnas_todas_semanas].mean(axis=1)
            else:
                # Para el TPI/CV también calculamos el promedio histórico redondeado a 2 decimales
                if 'Promedio' in df_actualizado.columns:
                    df_actualizado['Promedio'] = df_actualizado[columnas_todas_semanas].mean(axis=1).round(2)

            # --- CONVERSIÓN VISUAL FINAL (DE MINUTOS A FORMATO HH:MM) ---
            if columna_datos_a_extraer != 'TPI_Global':
                columnas_a_formatear = columnas_todas_semanas.copy()
                if 'Tiempo Acumulado' in df_actualizado.columns: columnas_a_formatear.append('Tiempo Acumulado')
                if 'Promedio' in df_actualizado.columns: columnas_a_formatear.append('Promedio')
                
                for col_fmt in columnas_a_formatear:
                    df_actualizado[col_fmt] = df_actualizado[col_fmt].apply(to_hhmm_display)

            # Guardamos la pestaña actualizada
            dict_dfs_actualizados[nombre_hoja_original] = df_actualizado.drop(columns=['MatchKey'], errors='ignore')

        else:
            dict_dfs_actualizados[nombre_hoja_original] = dict_dfs_originales[nombre_hoja_original]

    # --- INYECCIÓN DEL KPI DE ADHERENCIA EN EL EXCEL ---
    df_kpi_excel = df_semana_actual[['Deportista', 'TPI_Global', 'TPI_Natacion', 'TPI_Ciclismo', 'TPI_Trote', 'Es_Completo']].copy()
    
    for col_tpi in ['TPI_Global', 'TPI_Natacion', 'TPI_Ciclismo', 'TPI_Trote']:
        df_kpi_excel[col_tpi] = df_kpi_excel[col_tpi].apply(lambda x: f"{x:.1f}%" if pd.notna(x) else "0.0%")
        
    df_kpi_excel['Es_Completo'] = df_kpi_excel['Es_Completo'].apply(lambda x: "Sí" if x else "No")
    
    df_kpi_excel = df_kpi_excel.rename(columns={
        'Deportista': 'Nombre del Deportista',
        'TPI_Global': 'Adherencia Global',
        'TPI_Natacion': 'Adherencia Natación',
        'TPI_Ciclismo': 'Adherencia Ciclismo',
        'TPI_Trote': 'Adherencia Trote',
        'Es_Completo': f'Completó Plan ({etiqueta_semana})'
    })

    # --- NUEVA EXPANSIÓN: CREACIÓN DE LA HOJA CRUDA DE LA SEMANA ---
    df_hoja_semana = df_semana_actual[['Deportista', 'N_Mins_Real', 'B_Mins_Real', 'R_Mins_Real', 'T_Mins_Real']].copy()
    for col_tiempo in ['N_Mins_Real', 'B_Mins_Real', 'R_Mins_Real', 'T_Mins_Real']:
        df_hoja_semana[col_tiempo] = df_hoja_semana[col_tiempo].apply(to_hhmm_display)
        
    df_hoja_semana = df_hoja_semana.rename(columns={
        'Deportista': 'Nombre del Deportista',
        'N_Mins_Real': 'Natación',
        'B_Mins_Real': 'Bicicleta',
        'R_Mins_Real': 'Trote',
        'T_Mins_Real': 'Tiempo Total'
    })

    # --- ORDENAMIENTO ESTRICTO DE PESTAÑAS FINALES (JERARQUÍA Y DESCENDENTE) ---
    dict_final_ordenado = {}
    
    # Insertamos temporalmente las nuevas hojas para el escaneo
    dict_dfs_actualizados['KPI_Adherencia_TPI'] = df_kpi_excel
    dict_dfs_actualizados[etiqueta_semana] = df_hoja_semana
    
    # Jerarquía exacta solicitada
    orden_jerarquia = [
        'NUMEROS DE SEMANA', 'NUMERO DE SEMANA', 'SEMANAS', 'CALENDARIO',
        'TIEMPO TOTAL', 
        'NATACION', 'NATACIÓN', 
        'BICICLETA', 'CICLISMO', 
        'TROTE', 'RUNNING', 
        'CV', 
        'HISTORICO_TPI', 'HISTÓRICO_TPI', 'HISTORICOTPI',
        'KPI_ADHERENCIA_TPI'
    ]
    
    claves_base_ordenadas = []
    claves_semanas = []
    claves_otras = []
    
    for clave_real in dict_dfs_actualizados.keys():
        clave_limpia = clean_string(clave_real)
        
        # Aislar las hojas históricas de las semanas (Sem 08, Sem 07...)
        if str(clave_real).strip().upper().startswith('SEM '):
            claves_semanas.append(clave_real)
        else:
            if clave_limpia in orden_jerarquia:
                claves_base_ordenadas.append((orden_jerarquia.index(clave_limpia), clave_real))
            else:
                claves_otras.append(clave_real)
                
    # Ordenar las bases por la jerarquía solicitada
    claves_base_ordenadas.sort(key=lambda x: x[0])
    
    # Ordenar las semanas cronológicamente de forma descendente
    claves_semanas.sort(reverse=True)
    
    # Ensamblar el Excel Final
    for _, clave_real in claves_base_ordenadas:
        dict_final_ordenado[clave_real] = dict_dfs_actualizados[clave_real]
        
    for clave_semana in claves_semanas:
        dict_final_ordenado[clave_semana] = dict_dfs_actualizados[clave_semana]
        
    for clave_otra in claves_otras:
        dict_final_ordenado[clave_otra] = dict_dfs_actualizados[clave_otra]

    return dict_final_ordenado

# =============================================================================
# FIN DE SECCIÓN 5
# =============================================================================

# *****************************************************************************
# SECCIÓN 6: ORQUESTADOR DE ENTREGABLES (GRÁFICOS, COLORES Y REPORTES)
# *****************************************************************************
# Esta versión mantiene la estructura vertical extendida. Incluye el candado
# estricto de columnas para forzar a Word a respetar los márgenes, el centrado
# absoluto de tablas y el filtro de elegibilidad para TPI.
# NUEVO: Integración del Velocímetro Global del Equipo en la Portada (Dashboard Dual).
# NUEVO: Inyección de membretes corporativos en los ENCABEZADOS (Headers) de Word.
# NUEVO: Motor Narrativo Nivel 2 — diff_hist por disciplina, destino grupal/individual.

import matplotlib.pyplot as plt
import numpy as np
import io
import zipfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def ajustar_anchos_y_centrar_tabla(tabla, anchos):
    """
    Candado de formato: Desactiva el autoajuste, centra la tabla y fuerza
    el ancho desde la raíz de la columna para evitar desbordes en los márgenes.
    """
    tabla.autofit = False
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 1. Fijar el ancho de la columna desde la estructura principal
    for i, col in enumerate(tabla.columns):
        if i < len(anchos):
            col.width = anchos[i]
            
    # 2. Reforzar el ancho celda por celda para someter a Word
    for row in tabla.rows:
        for idx, width in enumerate(anchos):
            if idx < len(row.cells):
                row.cells[idx].width = width

def generar_velocimetro_tpi(porcentaje):
    """
    Dibuja un gráfico de velocímetro (Gauge Chart) semicircular para el TPI.
    """
    fig, ax = plt.subplots(figsize=(4, 2), subplot_kw={'projection': 'polar'})
    
    colors = ['#FF4C4C', '#FFD700', '#4CAF50']
    theta = np.linspace(0, np.pi, 100)
    
    ax.fill_between(np.linspace(np.pi, np.pi - (0.7 * np.pi), 50), 0.6, 1, color=colors[0], alpha=0.6)
    ax.fill_between(np.linspace(np.pi - (0.7 * np.pi), np.pi - (0.9 * np.pi), 50), 0.6, 1, color=colors[1], alpha=0.6)
    ax.fill_between(np.linspace(np.pi - (0.9 * np.pi), 0, 50), 0.6, 1, color=colors[2], alpha=0.6)
    
    p_grafico = min(porcentaje, 115)
    angulo_aguja = np.pi * (1 - (p_grafico / 115))
    
    ax.plot([angulo_aguja, angulo_aguja], [0, 0.9], color='black', linewidth=3, solid_capstyle='round')
    ax.plot(angulo_aguja, 0.9, marker='^', color='black', markersize=8)
    ax.plot(0, 0, marker='o', color='black', markersize=10)
    
    ax.set_ylim(0, 1)
    ax.set_yticks([])
    ax.set_xticks(np.pi * np.array([1, 0.39, 0.21, 0]))
    ax.set_xticklabels(['0%', '70%', '90%', '115%'], fontsize=10, weight='bold')
    ax.spines['polar'].set_visible(False)
    
    plt.text(0.5, 0.1, f"{porcentaje:.1f}%", transform=ax.transAxes, fontsize=20, weight='bold', ha='center', va='center')
             
    buf = io.BytesIO()
    plt.savefig(buf, format='png', transparent=True, bbox_inches='tight', dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf

def generar_grafico_distribucion(mins_natacion, mins_ciclismo, mins_trote):
    """
    Dibuja un gráfico de anillo para mostrar la distribución de disciplinas.
    Con textos en NEGRO y ajustados para legibilidad total en el Word.
    """
    fig, ax = plt.subplots(figsize=(4, 4))
    
    sizes = [mins_natacion, mins_ciclismo, mins_trote]
    labels = ['Natación', 'Ciclismo', 'Trote']
    colors = ['#1f77b4', '#ff7f0e', '#2ca02c']
    
    if sum(sizes) == 0:
        sizes = [1, 1, 1]
        labels = ['Sin Datos', 'Sin Datos', 'Sin Datos']
        
    wedges, texts, autotexts = ax.pie(
        sizes, 
        labels=labels, 
        autopct='%1.1f%%', 
        startangle=90, 
        colors=colors, 
        wedgeprops=dict(width=0.4, edgecolor='w'), 
        pctdistance=0.75,
        textprops={'color': 'black', 'weight': 'bold', 'size': 11}
    )
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', transparent=True, bbox_inches='tight', dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf

def generar_entregables_separados(df_semanal_procesado, dict_maestro_actualizado, etiqueta_semana):
    
    # =========================================================================
    # FASE 2: PRE-CÁLCULOS MATEMÁTICOS PARA ANÁLISIS COMPARATIVO
    # =========================================================================
    promedio_equipo = {
        'Total': df_semanal_procesado['T_Mins_Real'].mean(),
        'Natacion': df_semanal_procesado['N_Mins_Real'].mean(),
        'Ciclismo': df_semanal_procesado['B_Mins_Real'].mean(),
        'Trote': df_semanal_procesado['R_Mins_Real'].mean()
    }

    def obtener_media_historica(deportista_matchkey, hoja_maestro):
        # 1. Búsqueda dinámica de la hoja para superar diferencias de mayúsculas/tildes y sinónimos
        df_hoja = None
        hoja_limpia = clean_string(hoja_maestro)
        
        # Diccionario visual de sinónimos para asegurar cruce exacto
        sinonimos_bici = ['BICICLETA', 'CICLISMO']
        sinonimos_trote = ['TROTE', 'RUNNING']
        
        for key_real in dict_maestro_actualizado.keys():
            key_limpia = clean_string(key_real)
            
            if key_limpia == hoja_limpia:
                df_hoja = dict_maestro_actualizado[key_real]
                break
            elif hoja_limpia in sinonimos_bici and key_limpia in sinonimos_bici:
                df_hoja = dict_maestro_actualizado[key_real]
                break
            elif hoja_limpia in sinonimos_trote and key_limpia in sinonimos_trote:
                df_hoja = dict_maestro_actualizado[key_real]
                break
                
        if df_hoja is None or df_hoja.empty:
            return 0
            
        if 'Promedio' not in df_hoja.columns:
            return 0
            
        # 2. Identificación flexible de la columna de identidad (MatchKey ya no existe aquí)
        col_identidad_maestro = 'Nombre' if 'Nombre' in df_hoja.columns else \
                               ('Deportista' if 'Deportista' in df_hoja.columns else df_hoja.columns[0])
                               
        # 3. Búsqueda del atleta homologando la identidad en tiempo de ejecución
        mask_atleta = df_hoja[col_identidad_maestro].apply(clean_string) == deportista_matchkey
        fila_atleta = df_hoja[mask_atleta]
        
        if not fila_atleta.empty:
            return to_mins(fila_atleta['Promedio'].values[0])
        return 0

    def redactar_comparacion(minutos_reales, minutos_equipo, minutos_historicos):
        diff_equipo = minutos_reales - minutos_equipo
        if diff_equipo >= 0:
            txt_equipo = "MÁS"
        else:
            txt_equipo = "MENOS"
            
        val_equipo = to_hhmm_display(abs(diff_equipo))
        
        diff_hist = minutos_reales - minutos_historicos
        if diff_hist >= 0:
            txt_hist = "MÁS"
        else:
            txt_hist = "MENOS"
            
        val_hist = to_hhmm_display(abs(diff_hist))
        
        texto_final = f"Rendiste {val_equipo} {txt_equipo} que el promedio del equipo. Vs Tu Media Histórica: {val_hist} {txt_hist}."
        return texto_final

    # -------------------------------------------------------------------------
    # NUEVO: HELPER PARA CALCULAR DIFF_HIST POR DISCIPLINA ESPECÍFICA
    # -------------------------------------------------------------------------
    def calcular_diff_hist(fila, matchkey, hoja_maestro, col_real):
        """
        Calcula la diferencia en minutos entre el valor real de esta semana
        y la media histórica del atleta para una disciplina específica.
        Devuelve None si no hay historial previo (primer registro real).
        """
        promedio_mins = obtener_media_historica(matchkey, hoja_maestro)
        
        # Si el promedio histórico es 0, consideramos que es primer registro
        if promedio_mins == 0:
            return None
            
        return fila[col_real] - promedio_mins

    # =========================================================================
    # 6.1: GENERACIÓN DEL EXCEL MAESTRO ACTUALIZADO
    # =========================================================================
    buffer_excel = io.BytesIO()
    with pd.ExcelWriter(buffer_excel, engine='xlsxwriter') as writer:
        for nombre_hoja, df_hoja in dict_maestro_actualizado.items():
            df_hoja.to_excel(writer, sheet_name=nombre_hoja, index=False)
    buffer_excel.seek(0)
    
    # =========================================================================
    # 6.2: GENERACIÓN DEL REPORTE GRUPAL DEL CLUB (WORD AVANZADO)
    # =========================================================================
    doc_grupal = Document()
    
    # FORZAR CALIBRI EN TODO EL DOCUMENTO GRUPAL
    style_normal = doc_grupal.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)

    # --- NUEVO: MEMBRETE EN EL ENCABEZADO (HEADER) - REPORTE GRUPAL ---
    try:
        header_grupal = doc_grupal.sections[0].header
        p_head_g = header_grupal.paragraphs[0]
        p_head_g.alignment = WD_ALIGN_PARAGRAPH.RIGHT # Alineado a la esquina derecha
        p_head_g.add_run().add_picture("Tym Logo.jpg", width=Inches(0.8)) # Tamaño corporativo (0.8)
    except Exception:
        pass 
    # ------------------------------------------------------------------

    num_semana = str(etiqueta_semana).replace('Sem ', '')
    
    # --- PÁGINA 1: PORTADA, INTRODUCCIÓN Y DATOS GENERALES ---
    doc_grupal.add_heading(f"Reporte Semanal Club Tym Triatlón \nSemana {num_semana}", 0)
    doc_grupal.add_paragraph("") 
    
    doc_grupal.add_heading("1. Introducción General", level=1)
    p_intro = doc_grupal.add_paragraph()
    run_intro = p_intro.add_run("[ESPACIO PARA EDICIÓN MANUAL: Inserta aquí tu introducción de 4 líneas]")
    run_intro.bold = True
    run_intro.font.color.rgb = RGBColor(255, 0, 0)
    doc_grupal.add_paragraph("") 
    
    doc_grupal.add_heading("2. Datos Generales de la Semana", level=1)
    df_activos = df_semanal_procesado[df_semanal_procesado['T_Mins_Real'] > 0]
    mins_n = df_activos['N_Mins_Real'].sum()
    mins_b = df_activos['B_Mins_Real'].sum()
    mins_r = df_activos['R_Mins_Real'].sum()
    mins_totales = df_activos['T_Mins_Real'].sum()
    
    # Cálculo del nuevo KPI Global del Equipo
    tpi_promedio_equipo = df_activos['TPI_Global'].mean() if not df_activos.empty else 0.0
    
    p_res = doc_grupal.add_paragraph()
    p_res.add_run(f"Atletas activos esta semana: {len(df_activos)}\n").bold = True
    p_res.add_run(f"Horas de entrenamiento total del Equipo: {to_hhmm_display(mins_totales)}\n").bold = True
    p_res.add_run(f"Cumplimiento Promedio del Equipo (TPI): {tpi_promedio_equipo:.1f}%\n\n").bold = True
    p_res.add_run(f"Distribución por disciplina:\n")
    p_res.add_run(f"• Natación: {to_hhmm_display(mins_n)}\n")
    p_res.add_run(f"• Ciclismo: {to_hhmm_display(mins_b)}\n")
    p_res.add_run(f"• Trote: {to_hhmm_display(mins_r)}")
    doc_grupal.add_paragraph("") 
    
    # Dashboard Dual: Tabla invisible para gráficos lado a lado
    doc_grupal.add_heading("Visión Gráfica del Equipo", level=2)
    
    tabla_graficos = doc_grupal.add_table(rows=1, cols=2)
    # Ancho total = 5.6 pulgadas (Encaja perfecto en los márgenes de Word)
    ajustar_anchos_y_centrar_tabla(tabla_graficos, [Inches(2.8), Inches(2.8)])
    
    celda_anillo = tabla_graficos.rows[0].cells[0]
    celda_tpi = tabla_graficos.rows[0].cells[1]
    
    # Gráfico 1: Anillo de Disciplinas (Izquierda)
    img_dist = generar_grafico_distribucion(mins_n, mins_b, mins_r)
    p_anillo = celda_anillo.paragraphs[0]
    p_anillo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_anillo.add_run("Distribución de Horas\n").bold = True
    p_anillo.add_run().add_picture(img_dist, width=Inches(2.7))
    
    # Gráfico 2: Velocímetro Global (Derecha)
    img_tpi_global = generar_velocimetro_tpi(tpi_promedio_equipo)
    p_tpi = celda_tpi.paragraphs[0]
    p_tpi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tpi.add_run("Adherencia al Plan (TPI)\n").bold = True
    p_tpi.add_run().add_picture(img_tpi_global, width=Inches(2.7))
    
    doc_grupal.add_page_break() 

    # --- PÁGINA 2: TOP 5 COMPLETOS ---
    doc_grupal.add_heading("3. TOP 5 TRIATLETAS COMPLETOS", level=1)
    doc_grupal.add_paragraph("(Entrenamiento registrado en las 3 disciplinas)")
    
    tabla_completos = doc_grupal.add_table(rows=1, cols=6)
    tabla_completos.style = 'Light Grid Accent 1'
    
    # Ancho total = 5.8 pulgadas (Garantiza no salirse de la hoja)
    ajustar_anchos_y_centrar_tabla(tabla_completos, [Inches(0.4), Inches(2.2), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8)])
    
    cc = tabla_completos.rows[0].cells
    cc[0].text = '#'
    cc[1].text = 'Deportista'
    cc[2].text = 'Total'
    cc[3].text = 'Nat.'
    cc[4].text = 'Bici'
    cc[5].text = 'Trote'
    
    df_completos = df_semanal_procesado[df_semanal_procesado['Es_Completo'] == True]
    df_top_completos = df_completos.sort_values(by='T_Mins_Real', ascending=False).head(5)
    
    for pos, (_, fila) in enumerate(df_top_completos.iterrows(), 1):
        rc = tabla_completos.add_row().cells
        rc[0].text = str(pos)
        rc[1].text = str(fila['Deportista'])
        rc[2].text = to_hhmm_display(fila['T_Mins_Real'])
        rc[3].text = to_hhmm_display(fila['N_Mins_Real'])
        rc[4].text = to_hhmm_display(fila['B_Mins_Real'])
        rc[5].text = to_hhmm_display(fila['R_Mins_Real'])
        
    doc_grupal.add_paragraph("") 
        
    doc_grupal.add_heading("Comentarios del Grupo de Élite (Completos):", level=2)
    # PUNTO 3: diff_hist del tiempo total, destino grupal
    for pos, (_, fila) in enumerate(df_top_completos.iterrows(), 1):
        mk = fila['MatchKey']
        diff = calcular_diff_hist(fila, mk, 'TIEMPO TOTAL', 'T_Mins_Real')
        comentario = generar_comentario(fila, 'Completos', pos,
                                        diff_hist_mins=diff, destino='grupal')
        doc_grupal.add_paragraph(f"🏅 {comentario}")
        
    doc_grupal.add_page_break() 

    # --- PÁGINA 3: TOP 15 ADHERENCIA AL PLAN ---
    doc_grupal.add_heading("📈 4. TOP 15 ADHERENCIA AL PLAN (TPI GLOBAL)", level=1)
    doc_grupal.add_paragraph("") 
    
    # Filtro estricto para Adherencia
    def es_elegible_para_podio_tpi(row):
        if row.get('Natacion_Plan_Hrs', 0) > 0 and row.get('N_Mins_Real', 0) == 0: return False
        if row.get('Ciclismo_Plan_Hrs', 0) > 0 and row.get('B_Mins_Real', 0) == 0: return False
        if row.get('Trote_Plan_Hrs', 0) > 0 and row.get('R_Mins_Real', 0) == 0: return False
        if row.get('T_Mins_Real', 0) == 0: return False
        return True
        
    df_elegibles_tpi = df_semanal_procesado[df_semanal_procesado.apply(es_elegible_para_podio_tpi, axis=1)]
    df_ranking_tpi = df_elegibles_tpi.sort_values(by='TPI_Global', ascending=False).head(15)
    
    tabla_tpi = doc_grupal.add_table(rows=1, cols=3)
    tabla_tpi.style = 'Light Grid Accent 1'
    ajustar_anchos_y_centrar_tabla(tabla_tpi, [Inches(0.5), Inches(3.5), Inches(1.5)])
    
    c_tpi = tabla_tpi.rows[0].cells
    c_tpi[0].text = 'Pos.'
    c_tpi[1].text = 'Deportista'
    c_tpi[2].text = 'TPI Global %'
    
    for pos, (_, fila) in enumerate(df_ranking_tpi.iterrows(), 1):
        rc = tabla_tpi.add_row().cells
        rc[0].text = str(pos)
        rc[1].text = str(fila['Deportista'])
        rc[2].text = f"{fila['TPI_Global']:.1f}%"
        
    doc_grupal.add_paragraph("") 
        
    doc_grupal.add_heading("Comentarios del Podio (Adherencia):", level=2)
    medallas = {1: "🥇", 2: "🥈", 3: "🥉"}
    # PUNTO 4: diff_hist del tiempo total, destino grupal
    for pos, (_, fila) in enumerate(df_ranking_tpi.head(3).iterrows(), 1):
        mk = fila['MatchKey']
        diff = calcular_diff_hist(fila, mk, 'TIEMPO TOTAL', 'T_Mins_Real')
        comentario = generar_comentario(fila, 'TPI', pos,
                                        diff_hist_mins=diff, destino='grupal')
        doc_grupal.add_paragraph(f"{medallas.get(pos, '')} {comentario}")
        
    doc_grupal.add_page_break() 

    # --- PÁGINAS 4 a 7: TOP 15 POR TIEMPO Y DISCIPLINAS ---
    # PUNTO 2: bloques_tops ahora incluye hoja_hist para calcular diff por disciplina
    bloques_tops = [
        ("⏱️ 5. TOP 15 TIEMPO TOTAL", 'T_Mins_Real', 'General',   'TIEMPO TOTAL'),
        ("🏊‍♂️ 6. TOP 15 NATACIÓN",   'N_Mins_Real', 'Natación',  'NATACION'),
        ("🚴‍♂️ 7. TOP 15 CICLISMO",   'B_Mins_Real', 'Bicicleta', 'BICICLETA'),
        ("🏃‍♂️ 8. TOP 15 TROTE",      'R_Mins_Real', 'Trote',     'TROTE')
    ]
    
    for titulo, columna, categoria_frase, hoja_hist in bloques_tops:
        doc_grupal.add_heading(titulo, level=1)
        doc_grupal.add_paragraph("") 
        
        tabla_disc = doc_grupal.add_table(rows=1, cols=3)
        tabla_disc.style = 'Light Grid Accent 1'
        ajustar_anchos_y_centrar_tabla(tabla_disc, [Inches(0.5), Inches(3.5), Inches(1.5)])
        
        cd = tabla_disc.rows[0].cells
        cd[0].text = 'Pos.'
        cd[1].text = 'Deportista'
        cd[2].text = 'Tiempo'
        
        df_disc = df_semanal_procesado[df_semanal_procesado[columna] > 0].sort_values(by=columna, ascending=False).head(15)
        for pos, (_, fila) in enumerate(df_disc.iterrows(), 1):
            rc = tabla_disc.add_row().cells
            rc[0].text = str(pos)
            rc[1].text = str(fila['Deportista'])
            rc[2].text = to_hhmm_display(fila[columna])
            
        doc_grupal.add_paragraph("") 
            
        doc_grupal.add_heading(f"Comentarios del Podio ({categoria_frase}):", level=2)
        # PUNTO 5: diff_hist calculado para la disciplina específica del ranking
        for pos, (_, fila) in enumerate(df_disc.head(3).iterrows(), 1):
            mk = fila['MatchKey']
            diff = calcular_diff_hist(fila, mk, hoja_hist, columna)
            comentario = generar_comentario(fila, categoria_frase, pos,
                                            diff_hist_mins=diff, destino='grupal')
            doc_grupal.add_paragraph(f"{medallas.get(pos, '')} {comentario}")
        
        doc_grupal.add_page_break()

    # --- PÁGINA 8: ESTRATEGIA Y CONCLUSIONES ---
    doc_grupal.add_heading("💡 9. Insights Estratégicos", level=1)
    doc_grupal.add_paragraph("") 
    
    p_ins = doc_grupal.add_paragraph()
    run_ins_1 = p_ins.add_run("Bici-dependencia Crónica: ")
    run_ins_1.bold = True
    p_ins.add_run("El ciclismo sigue siendo el \"hijo favorito\". Representa la mayor parte del tiempo total del club. Básicamente, si le quitamos las bicicletas a este equipo, el reporte semanal cabría en una servilleta.\n")
    
    run_ins_2 = p_ins.add_run("El \"Efecto Claudio\": ")
    run_ins_2.bold = True
    p_ins.add_run("Una lección magistral de cómo liderar el ciclismo y aun así bajar en el ranking general por el \"pequeño\" detalle de no tocar el agua. Recordatorio amistoso: el Triatlón, por definición, incluye nadar.\n")
    
    run_ins_3 = p_ins.add_run("Resiliencia en el Asfalto: ")
    run_ins_3.bold = True
    p_ins.add_run("Mientras el volumen general bajaba, el trote sacó la cara con registros notables de carrera a pie. Parece que algunos sí desayunaron ganas de correr.\n")
    
    doc_grupal.add_paragraph("") 
    
    doc_grupal.add_heading("10. Conclusiones Generales", level=1)
    p_cierre = doc_grupal.add_paragraph()
    run_cierre = p_cierre.add_run("[ESPACIO PARA EDICIÓN MANUAL: Inserta aquí tus conclusiones finales y el cierre de la semana]")
    run_cierre.bold = True
    run_cierre.font.color.rgb = RGBColor(255, 0, 0)

    buffer_word_grupal = io.BytesIO()
    doc_grupal.save(buffer_word_grupal)
    buffer_word_grupal.seek(0)

    # =========================================================================
    # 6.3: GENERACIÓN DE FICHAS INDIVIDUALES (ZIP) 
    # =========================================================================
    buffer_zip_fichas = io.BytesIO()
    
    with zipfile.ZipFile(buffer_zip_fichas, "a", zipfile.ZIP_DEFLATED) as archivo_zip:
        
        for index, row_atleta in df_semanal_procesado.iterrows():
            if row_atleta['T_Mins_Real'] > 0:
                doc_i = Document()
                
                style_ind = doc_i.styles['Normal']
                font_ind = style_ind.font
                font_ind.name = 'Calibri'
                font_ind.size = Pt(11)
                
                # --- NUEVO: MEMBRETE DOBLE EN EL ENCABEZADO - FICHAS INDIVIDUALES ---
                try:
                    header_i = doc_i.sections[0].header
                    # Creamos una tabla invisible en el encabezado para posicionar en los extremos
                    tabla_header = header_i.add_table(rows=1, cols=2, width=Inches(6.5))
                    
                    # Logo Metri KM a la Izquierda
                    celda_izq = tabla_header.cell(0, 0)
                    p_izq = celda_izq.paragraphs[0]
                    p_izq.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_izq.add_run().add_picture("logo_metrikm.png", width=Inches(0.8))
                    
                    # Logo TYM a la Derecha
                    celda_der = tabla_header.cell(0, 1)
                    p_der = celda_der.paragraphs[0]
                    p_der.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p_der.add_run().add_picture("Tym Logo.jpg", width=Inches(0.8))
                except Exception:
                    pass
                # --------------------------------------------------------------------
                
                doc_i.add_heading(f"Análisis de Rendimiento Personal: {row_atleta['Deportista']}", 0)
                
                p_sub = doc_i.add_paragraph(f"Semana de Entrenamiento: {num_semana}")
                p_sub.bold = True
                doc_i.add_paragraph("") 
                
                # --- BLOQUE 1: VELOCÍMETRO DE ADHERENCIA (TPI) ---
                doc_i.add_heading("CUMPLIMIENTO DE PLAN (TPI)", level=2)
                p_tpi_intro = doc_i.add_paragraph("Tu Índice de Adherencia Global esta semana:")
                p_tpi_intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                img_velocimetro = generar_velocimetro_tpi(row_atleta['TPI_Global'])
                para_img = doc_i.add_paragraph()
                para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para_img.add_run().add_picture(img_velocimetro, width=Inches(3.5))
                doc_i.add_paragraph("") 
                
                tabla_desglose = doc_i.add_table(rows=1, cols=4)
                tabla_desglose.style = 'Table Grid'
                ajustar_anchos_y_centrar_tabla(tabla_desglose, [Inches(1.5), Inches(1.2), Inches(1.2), Inches(1.2)])
                
                c_cab = tabla_desglose.rows[0].cells
                c_cab[0].text = 'Disciplina'
                c_cab[1].text = 'Real'
                c_cab[2].text = 'Meta'
                c_cab[3].text = 'TPI %'
                
                matriz = [
                    ('Natación', 'N_Mins_Real', 'Natacion_Plan_Hrs', 'TPI_Natacion'),
                    ('Ciclismo', 'B_Mins_Real', 'Ciclismo_Plan_Hrs', 'TPI_Ciclismo'),
                    ('Trote', 'R_Mins_Real', 'Trote_Plan_Hrs', 'TPI_Trote')
                ]
                
                for d, c_r, c_m, c_tpi in matriz:
                    rc = tabla_desglose.add_row().cells
                    rc[0].text = d
                    rc[1].text = to_hhmm_display(row_atleta[c_r])
                    rc[2].text = f"{row_atleta[c_m]:.1f}h"
                    
                    run_tpi = rc[3].paragraphs[0].add_run(f"{row_atleta[c_tpi]:.1f}%")
                    run_tpi.bold = True
                    
                    tpi_val = row_atleta[c_tpi]
                    if tpi_val < 70:
                        run_tpi.font.color.rgb = RGBColor(255, 0, 0)
                    elif tpi_val <= 90:
                        run_tpi.font.color.rgb = RGBColor(204, 153, 0)
                    else:
                        run_tpi.font.color.rgb = RGBColor(0, 153, 0)
                        
                doc_i.add_paragraph("") 

                # --- BLOQUE 2: ANÁLISIS HISTÓRICO COMPARATIVO ---
                mk = row_atleta['MatchKey']
                bloques_analisis = [
                    ('TIEMPO TOTAL', 'T_Mins_Real', 'Total', 'TIEMPO TOTAL'),
                    ('NATACIÓN', 'N_Mins_Real', 'Natacion', 'NATACION'),
                    ('CICLISMO', 'B_Mins_Real', 'Ciclismo', 'BICICLETA'),
                    ('TROTE', 'R_Mins_Real', 'Trote', 'TROTE')
                ]
                
                for titulo_bloque, col_real, llave_equipo, hoja_historia in bloques_analisis:
                    doc_i.add_heading(titulo_bloque, level=2)
                    p_dato = doc_i.add_paragraph()
                    
                    run_volumen = p_dato.add_run(f"Volumen actual: {to_hhmm_display(row_atleta[col_real])} ")
                    run_volumen.bold = True
                    
                    hist_mins = obtener_media_historica(mk, hoja_historia)
                    texto_comp = redactar_comparacion(row_atleta[col_real], promedio_equipo[llave_equipo], hist_mins)
                    p_dato.add_run(texto_comp)
                    
                doc_i.add_paragraph("") 

                # --- BLOQUE 3: EVALUACIÓN TÉCNICA (MOTOR NARRATIVO NIVEL 2) ---
                doc_i.add_heading("Evaluación Técnica", level=2)

                # PUNTO 6: diff_hist del tiempo total para el comentario general
                mk = row_atleta['MatchKey']
                diff_total = calcular_diff_hist(row_atleta, mk, 'TIEMPO TOTAL', 'T_Mins_Real')
                comentario_general = generar_comentario(
                    row_atleta, 'General', 1,
                    diff_hist_mins=diff_total,
                    destino='individual'
                )
                doc_i.add_paragraph(comentario_general)

                # PUNTO 6: Comentario específico por cada disciplina activa
                # con diff_hist calculado para esa disciplina en particular
                disc_ficha = [
                    ('Natación',  'N_Mins_Real', 'NATACION'),
                    ('Bicicleta', 'B_Mins_Real', 'BICICLETA'),
                    ('Trote',     'R_Mins_Real', 'TROTE'),
                ]
                for nombre_disc, col_disc, hoja_disc in disc_ficha:
                    if row_atleta[col_disc] > 0:
                        diff_disc = calcular_diff_hist(
                            row_atleta, mk, hoja_disc, col_disc)
                        comentario_disc = generar_comentario(
                            row_atleta, nombre_disc, 1,
                            diff_hist_mins=diff_disc,
                            destino='individual'
                        )
                        doc_i.add_paragraph(comentario_disc)
                
                doc_i.add_paragraph("\n──────────────────────────────────────────────────")
                doc_i.add_paragraph("Generado por Plataforma Metri KM - TYM")
                
                buffer_word_individual = io.BytesIO()
                doc_i.save(buffer_word_individual)
                
                nombre_archivo = f"Reporte_{clean_string(row_atleta['Deportista'])}.docx"
                archivo_zip.writestr(nombre_archivo, buffer_word_individual.getvalue())

    buffer_zip_fichas.seek(0)
    
    return {
        'excel_maestro': buffer_excel,
        'word_grupal': buffer_word_grupal,
        'zip_fichas': buffer_zip_fichas
    }

# =============================================================================
# FIN DE SECCIÓN 6
# =============================================================================
    
# *****************************************************************************
# SECCIÓN 7: INTERFAZ DE USUARIO, ORQUESTACIÓN Y CONSOLA DE DEBUG
# *****************************************************************************

# -----------------------------------------------------------------------------
# 7.1: CONFIGURACIÓN VISUAL DEL PANEL CENTRAL Y SIDEBAR
# -----------------------------------------------------------------------------
# Configuración de la pestaña del navegador
st.set_page_config(
    page_title="Metri KM - TYM",
    page_icon="⏱️",
    layout="wide"
)

# --- Cabecera Principal (Logos y Título) ---
col_izq, col_centro, col_der = st.columns([1, 6, 1])

with col_izq:
    try:
        st.image("logo_metrikm.png", use_container_width=True)
    except:
        st.info("🖼️ Metri KM")

with col_centro:
    st.title("Metri KM - TYM")
    st.markdown("### Plataforma de Inteligencia Deportiva y Adherencia al Plan")

with col_der:
    try:
        st.image("Tym Logo.jpg", use_container_width=True)
    except:
        st.info("🖼️ TYM")

# --- Panel Lateral (Sidebar) ---
with st.sidebar:
    # Logo superior en el menú de controles
    try:
        st.image("logo_metrikm.png", use_container_width=True)
    except:
        pass
        
    st.header("⚙️ 1. Entradas de Sistema")
    st.markdown("Carga los archivos obligatorios para iniciar el procesamiento.")
    
    archivo_maestro = st.file_uploader("A. Maestro Histórico (Excel)", type=["xlsx", "xls"])
    archivo_strava = st.file_uploader("B. Strava Semanal (Excel)", type=["xlsx", "xls"])
    archivo_plan = st.file_uploader("C. Plan Individual (Opcional)", type=["xlsx", "xls"])
    
    st.divider()
    
    st.header("🎯 2. Metas Globales (Fallback)")
    st.markdown("Aplicables si un atleta no tiene Plan Individual.")
    metas_globales_sidebar = {
        'N_H': st.number_input("Natación - Horas", min_value=0.0, value=3.0, step=0.5),
        'N_S': st.number_input("Natación - Sesiones", min_value=1, value=3, step=1),
        'B_H': st.number_input("Ciclismo - Horas", min_value=0.0, value=5.0, step=0.5),
        'B_S': st.number_input("Ciclismo - Sesiones", min_value=1, value=3, step=1),
        'T_H': st.number_input("Trote - Horas", min_value=0.0, value=3.0, step=0.5),
        'T_S': st.number_input("Trote - Sesiones", min_value=1, value=3, step=1)
    }
    
    st.divider()
    
    st.header("🏷️ 3. Etiqueta Temporal")
    etiqueta_semana_input = st.text_input("Etiqueta a generar en el Maestro", "Sem 08")

# -----------------------------------------------------------------------------
# 7.2: INICIALIZACIÓN DEL ESTADO DE SESIÓN (SESSION STATE)
# -----------------------------------------------------------------------------
# Esto es vital para que los tres botones de descarga no desaparezcan 
# de la pantalla al hacer clic en uno de ellos.
if 'diccionario_entregables' not in st.session_state:
    st.session_state['diccionario_entregables'] = None

# -----------------------------------------------------------------------------
# 7.3: EJECUCIÓN PRINCIPAL (BOTÓN DE PROCESAMIENTO)
# -----------------------------------------------------------------------------
if st.button("🚀 PROCESAR EXCEL Y GENERAR ENTREGABLES", use_container_width=True):
    
    # Verificación de seguridad: No avanzar sin los archivos vitales
    if archivo_maestro is not None and archivo_strava is not None:
        
        with st.spinner("Purificando datos, calculando Adherencia TYM y generando gráficas..."):
            try:
                # --- FASE 1: EXTRACCIÓN (Sección 2) ---
                df_datos_reales = procesar_strava_excel(archivo_strava)
                df_datos_plan = procesar_plan_individual(archivo_plan)
                
                # --- FASE 2: CÁLCULO DE KPI (Sección 4) ---
                df_calculado_kpi = calcular_kpis_tym(df_datos_reales, df_datos_plan, metas_globales_sidebar)
                
                # --- FASE 3: CONSOLA DE AUDITORÍA (DEBUG VISUAL) ---
                # Mostrar en pantalla los resultados antes de empaquetarlos para validar precisión
                with st.expander("🕵️‍♂️ CONSOLA DE AUDITORÍA (VERIFICAR ANTES DE DESCARGAR)", expanded=True):
                    st.markdown("**1. Lectura Pura de Strava (Mins Reales Extraídos):**")
                    st.dataframe(df_calculado_kpi[['Deportista', 'N_Mins_Real', 'B_Mins_Real', 'R_Mins_Real', 'T_Mins_Real']].head(10))
                    
                    st.markdown("**2. Cálculo de KPIs (Adherencia):**")
                    st.dataframe(df_calculado_kpi[['Deportista', 'TPI_Natacion', 'TPI_Ciclismo', 'TPI_Trote', 'TPI_Global', 'Es_Completo']].head(10))

                # --- FASE 4: PERSISTENCIA DEL MAESTRO (Sección 5) ---
                diccionario_maestro_original = pd.read_excel(archivo_maestro, sheet_name=None)
                diccionario_maestro_actualizado = actualizar_maestro_tym(
                    diccionario_maestro_original, 
                    df_calculado_kpi, 
                    etiqueta_semana_input
                )
                
                # --- FASE 5: GENERACIÓN DE ENTREGABLES SEPARADOS (Sección 6) ---
                st.session_state['diccionario_entregables'] = generar_entregables_separados(
                    df_calculado_kpi, 
                    diccionario_maestro_actualizado, 
                    etiqueta_semana_input
                )
                
                st.success("✅ ¡Procesamiento completado con éxito! Gráficos generados. Revisa la consola y descarga tus archivos.")
                
            except Exception as error_critico:
                st.error(f"❌ Error crítico durante el procesamiento: {str(error_critico)}")
    else:
        st.warning("⚠️ Debes cargar el Maestro Histórico y el Excel de Strava en el panel lateral para iniciar.")

# -----------------------------------------------------------------------------
# 7.4: DESPLIEGUE DE LOS TRES BOTONES DE DESCARGA INDEPENDIENTES
# -----------------------------------------------------------------------------
if st.session_state['diccionario_entregables'] is not None:
    st.markdown("### 📥 Descarga de Archivos Procesados")
    st.markdown("Selecciona el entregable que deseas descargar:")
    
    columna_excel, columna_word, columna_zip = st.columns(3)
    
    with columna_excel:
        st.download_button(
            label="📊 1. Descargar Maestro Actualizado (.xlsx)",
            data=st.session_state['diccionario_entregables']['excel_maestro'],
            file_name=f"01_Estadisticas_TYM_{etiqueta_semana_input}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )
        
    with columna_word:
        st.download_button(
            label="📄 2. Descargar Reporte Grupal (.docx)",
            data=st.session_state['diccionario_entregables']['word_grupal'],
            file_name=f"02_Reporte_General_{etiqueta_semana_input}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
    with columna_zip:
        st.download_button(
            label="🗂️ 3. Descargar Fichas Individuales (.zip)",
            data=st.session_state['diccionario_entregables']['zip_fichas'],
            file_name=f"03_Fichas_Individuales_{etiqueta_semana_input}.zip",
            mime="application/zip",
            use_container_width=True
        )
